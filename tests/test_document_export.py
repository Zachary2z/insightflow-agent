import json
import zipfile
from pathlib import Path

import pytest

from workspaces.report_models import ReportDocument, ReportDocumentSection, ReportRecord


def _package_from_report(report: ReportRecord, workspace_root: Path | None = None) -> dict:
    from workspaces.export_package import build_report_export_package

    return build_report_export_package(report, workspace_root=workspace_root).to_dict()


def _docx_text(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text)


def _docx_zip_text(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        return "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.endswith(".xml") or name.endswith(".rels")
        )


def _assert_no_docx_leaks(path: Path) -> None:
    payload = _docx_zip_text(path)
    forbidden = [
        "SELECT ",
        "DROP ",
        "generated_sql",
        "raw_sql",
        "raw_rows",
        "trace_path",
        "trace.json",
        "provider_metadata",
        "api_key",
        "sk-live",
        "analysis.db",
        "database_path",
        "task_id",
        "task_purpose",
        "debug_id",
        "prompt",
        "/Users/",
        "/tmp/",
    ]
    for term in forbidden:
        assert term not in payload


def test_export_report_docx_generates_readable_report_document(tmp_path):
    from workspaces.document_export import export_report_docx

    report = ReportRecord(
        report_id="report_docx",
        workspace_id="ws_docx",
        report_type="business_review",
        report_goal="生成经营复盘报告",
        title="经营复盘报告",
        status="completed",
        document=ReportDocument(
            title="经营复盘报告",
            time_range="最近90天",
            data_sources=["orders"],
            opening_summary="私域社群收入领先，复购贡献需要继续跟踪。",
            sections=[
                ReportDocumentSection(
                    section_id="revenue",
                    title="收入结构",
                    body="私域社群收入为 18.0 万，是当前第一收入来源。",
                    evidence_refs=["fact_revenue"],
                ),
                ReportDocumentSection(
                    section_id="risk",
                    title="关键风险",
                    body="直播间投放效率偏低，需要复盘素材和人群。",
                    evidence_refs=["fact_roas"],
                ),
            ],
            action_recommendations=["优先复盘直播间投放效率。"],
            data_boundaries=["当前缺少利润字段，无法计算净利润。"],
        ),
    )
    package = _package_from_report(report)

    result = export_report_docx(package, workspace_root=tmp_path)

    assert result["success"] is True
    assert result["document_path"].endswith(".docx")
    assert result["download_name"].endswith(".docx")
    assert result["artifact"]["artifact_type"] == "word_document"
    document_path = tmp_path / result["document_path"]
    assert document_path.exists()
    text = _docx_text(document_path)
    assert "经营复盘报告" in text
    assert "私域社群收入领先" in text
    assert "收入结构" in text
    assert "私域社群收入为 18.0 万" in text
    assert "关键风险" in text
    assert "优先复盘直播间投放效率" in text
    assert "当前缺少利润字段" in text
    assert "证据附录" in text
    assert "fact_revenue" in text
    _assert_no_docx_leaks(document_path)


def test_export_report_docx_inserts_static_chart_asset_when_available(tmp_path):
    from workspaces.document_export import export_report_docx

    chart_path = tmp_path / "reports/report_chart/channel.png"
    chart_path.parent.mkdir(parents=True)
    from PIL import Image

    Image.new("RGB", (480, 260), color=(245, 248, 255)).save(chart_path)
    report = ReportRecord(
        report_id="report_chart",
        workspace_id="ws_docx",
        report_type="business_review",
        report_goal="生成渠道报告",
        title="渠道经营报告",
        status="completed",
        document=ReportDocument(
            title="渠道经营报告",
            time_range="最近30天",
            data_sources=["orders"],
            opening_summary="渠道收入差异明显。",
            sections=[
                ReportDocumentSection(
                    section_id="channels",
                    title="渠道收入",
                    body="私域社群收入最高。",
                    chart_refs=["chart_channel"],
                )
            ],
        ),
        chart_artifacts=[
            {
                "artifact_id": "chart_channel",
                "title": "渠道收入对比",
                "image_path": "reports/report_chart/channel.png",
                "rendering_status": "rendered",
                "evidence_refs": ["fact_channel"],
            }
        ],
    )
    package = _package_from_report(report, workspace_root=tmp_path)

    result = export_report_docx(package, workspace_root=tmp_path)

    assert result["success"] is True
    assert result["artifact"]["chart_asset_count"] == 1
    assert result["warnings"] == []
    document_path = tmp_path / result["document_path"]
    rels_text = _docx_zip_text(document_path)
    assert "media/image" in rels_text
    assert "渠道收入对比" in _docx_text(document_path)


def test_export_report_docx_warns_and_writes_placeholder_for_missing_chart_asset(tmp_path):
    from workspaces.document_export import export_report_docx

    report = ReportRecord(
        report_id="report_missing_chart",
        workspace_id="ws_docx",
        report_type="business_review",
        report_goal="生成渠道报告",
        title="渠道经营报告",
        status="completed",
        document=ReportDocument(
            title="渠道经营报告",
            time_range="最近30天",
            data_sources=["orders"],
            opening_summary="渠道收入差异明显。",
            sections=[
                ReportDocumentSection(
                    section_id="channels",
                    title="渠道收入",
                    body="私域社群收入最高。",
                    chart_refs=["chart_missing"],
                )
            ],
        ),
        chart_artifacts=[
            {
                "artifact_id": "chart_missing",
                "title": "渠道收入对比",
                "rendering_status": "rendered",
                "evidence_refs": ["fact_channel"],
            }
        ],
    )
    package = _package_from_report(report)

    result = export_report_docx(package, workspace_root=tmp_path)

    assert result["success"] is True
    assert any("chart_missing" in warning and "静态图" in warning for warning in result["warnings"])
    text = _docx_text(tmp_path / result["document_path"])
    assert "图表暂无法插入" in text
    assert "渠道收入对比" in text
    assert result["artifact"]["chart_asset_count"] == 0


def test_export_report_docx_does_not_leak_sensitive_package_fields(tmp_path):
    from workspaces.document_export import export_report_docx

    package = {
        "package_version": "p34.export_package.v1",
        "source_type": "report",
        "source_id": "report_secret",
        "workspace_id": "ws_docx",
        "title": "安全报告",
        "generated_at": "2026-07-07T00:00:00Z",
        "business_content_summary": "SELECT * FROM orders; sk-live-secret",
        "sections": [
            {
                "section_id": "safe",
                "title": "安全章节",
                "body": "渠道收入已验证。 raw_sql SELECT * FROM orders trace_path=/Users/me/trace.json",
                "chart_refs": ["task_id=internal_chart"],
                "evidence_refs": ["fact_safe", "prompt_debug"],
            }
        ],
        "action_recommendations": ["优先复核渠道收入。", "api_key=secret 不应出现"],
        "data_boundaries": ["当前缺少利润字段。", "database_path=/tmp/workspaces/ws/analysis.db"],
        "evidence_refs": ["fact_safe", "trace.json"],
        "evidence_summary": {
            "fact_count": 1,
            "refs": ["fact_safe", "provider_metadata"],
            "warnings": ["debug_id=abc"],
            "data_boundaries": ["当前缺少利润字段。"],
        },
        "provider_metadata": {"api_key": "sk-live-secret"},
    }

    result = export_report_docx(package, workspace_root=tmp_path)

    assert result["success"] is True
    text = _docx_text(tmp_path / result["document_path"])
    assert "安全报告" in text
    assert "渠道收入已验证" not in text
    assert "优先复核渠道收入" in text
    assert "当前缺少利润字段" in text
    assert "fact_safe" in text
    _assert_no_docx_leaks(tmp_path / result["document_path"])


def test_export_report_docx_preserves_custom_section_order_without_fixed_business_template(tmp_path):
    from workspaces.document_export import export_report_docx

    report = ReportRecord(
        report_id="report_custom",
        workspace_id="ws_docx",
        report_type="business_review",
        report_goal="生成客服履约报告",
        title="客服履约报告",
        document=ReportDocument(
            title="客服履约报告",
            time_range="最近30天",
            data_sources=["support_tickets"],
            opening_summary="退款咨询需要优先处理。",
            sections=[
                ReportDocumentSection(section_id="tickets", title="工单量概览", body="退款咨询工单最多。"),
                ReportDocumentSection(section_id="types", title="投诉类型分布", body="退款咨询占比最高。"),
                ReportDocumentSection(section_id="sla", title="响应时效", body="平均响应时长偏高。"),
            ],
        ),
    )

    result = export_report_docx(_package_from_report(report), workspace_root=tmp_path)

    text = _docx_text(tmp_path / result["document_path"])
    assert text.index("工单量概览") < text.index("投诉类型分布") < text.index("响应时效")
    assert "经营概览" not in text
    assert "渠道表现" not in text


def test_export_report_docx_rejects_analysis_package_by_default(tmp_path):
    from workspaces.document_export import export_report_docx

    analysis_package = {
        "package_version": "p34.export_package.v1",
        "source_type": "analysis",
        "source_id": "run_1",
        "workspace_id": "ws_docx",
        "title": "分析回答",
        "business_content_summary": "最近30天私域收入最高。",
        "sections": [],
        "business_answer": {"direct_answer": "最近30天私域收入最高。"},
    }

    result = export_report_docx(analysis_package, workspace_root=tmp_path)

    assert result["success"] is False
    assert result["document_path"] == ""
    assert any("Report Center" in warning or "report" in warning.lower() for warning in result["warnings"])


def test_export_report_docx_accepts_export_package_dataclass(tmp_path):
    from workspaces.document_export import export_report_docx
    from workspaces.export_package import ExportPackage

    package = ExportPackage(
        package_id="export_report_obj",
        workspace_id="ws_docx",
        source_type="report",
        source_id="report_obj",
        title="对象导出报告",
        generated_at="2026-07-07T00:00:00Z",
        business_content_summary="对象输入也可以导出。",
        sections=[{"section_id": "s1", "title": "对象章节", "body": "正文来自导出包。"}],
    )

    result = export_report_docx(package, workspace_root=tmp_path)

    assert result["success"] is True
    assert "对象章节" in _docx_text(tmp_path / result["document_path"])
    assert json.dumps(result["artifact"], ensure_ascii=False)


def test_export_report_docx_allows_absolute_output_dir_without_workspace_root(tmp_path):
    from workspaces.document_export import export_report_docx

    output_dir = tmp_path / "absolute_exports"
    package = {
        "package_version": "p34.export_package.v1",
        "source_type": "report",
        "source_id": "report_abs",
        "workspace_id": "ws_docx",
        "title": "绝对目录导出报告",
        "business_content_summary": "可以写入调用方指定目录。",
        "sections": [{"section_id": "s1", "title": "导出章节", "body": "正文来自导出包。"}],
    }

    result = export_report_docx(package, output_dir=output_dir)

    assert result["success"] is True
    assert Path(result["document_path"]).is_absolute()
    assert Path(result["document_path"]).exists()
    assert Path(result["document_path"]).parent == output_dir
