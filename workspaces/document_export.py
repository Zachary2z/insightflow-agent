from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
import re
from typing import Any
from urllib.parse import urlsplit

from workspaces.export_package import EXPORT_PACKAGE_VERSION, SOURCE_REPORT
from workspaces.models import utc_now_iso


DEFAULT_DOCUMENT_EXPORT_DIR = "exports/documents"
SUPPORTED_WORD_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tif", ".tiff"}


@dataclass
class DocumentExportPayload:
    workspace_id: str
    source_id: str
    title: str
    generated_at: str
    summary: str = ""
    time_range: str = ""
    data_sources: list[str] = field(default_factory=list)
    sections: list[dict[str, Any]] = field(default_factory=list)
    action_recommendations: list[str] = field(default_factory=list)
    data_boundaries: list[str] = field(default_factory=list)
    chart_artifacts: dict[str, dict[str, Any]] = field(default_factory=dict)
    static_assets: dict[str, dict[str, Any]] = field(default_factory=dict)
    evidence_refs: list[str] = field(default_factory=list)
    evidence_summary: dict[str, Any] = field(default_factory=dict)


def export_report_docx(
    export_package: Any,
    *,
    workspace_root: str | Path | None = None,
    output_dir: str | Path | None = None,
) -> dict[str, Any]:
    """Render an existing safe Report Center export package into a local Word document."""
    package = _to_dict(export_package)
    warnings: list[str] = []
    if not _is_supported_report_package(package):
        return _failed_result(
            "Word 导出当前只支持 Report Center 的 report export package；Analysis Workbench 轻量分析导出尚未启用。"
        )

    payload = _payload_from_package(package)
    if not payload.title:
        return _failed_result("导出包缺少安全的报告标题，无法生成 Word 文档。")

    target = _document_target(payload, workspace_root=workspace_root, output_dir=output_dir)
    if not target:
        return _failed_result("无法确定安全的 Word 文档输出目录。")

    try:
        from docx import Document
    except ModuleNotFoundError:
        return _failed_result("缺少 python-docx 依赖，无法生成 Word 文档。")

    target["absolute_path"].parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_document_styles(document)
    chart_asset_count = _render_document(
        document=document,
        payload=payload,
        workspace_root=workspace_root,
        warnings=warnings,
    )
    document.save(str(target["absolute_path"]))

    artifact = {
        "artifact_id": f"artifact_docx_{payload.source_id or _safe_filename(payload.title)}",
        "artifact_type": "word_document",
        "title": payload.title,
        "relative_path": target["relative_path"],
        "download_name": target["download_name"],
        "source": "document_export",
        "status": "completed",
        "created_at": utc_now_iso(),
        "chart_asset_count": chart_asset_count,
        "source_package_version": EXPORT_PACKAGE_VERSION,
    }
    return {
        "success": True,
        "document_path": target["relative_path"],
        "download_name": target["download_name"],
        "warnings": _unique_texts(warnings),
        "artifact": artifact,
    }


def _render_document(
    *,
    document: Any,
    payload: DocumentExportPayload,
    workspace_root: str | Path | None,
    warnings: list[str],
) -> int:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(payload.title)

    meta_items = [item for item in [payload.time_range, "、".join(payload.data_sources)] if item]
    if meta_items:
        paragraph = document.add_paragraph(style="InsightFlowMeta")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(" | ".join(meta_items))

    if payload.summary:
        document.add_heading("摘要", level=1)
        _add_body_text(document, payload.summary)

    chart_asset_count = 0
    for section_data in payload.sections:
        section_title = _safe_text(section_data.get("title"))
        section_body = _safe_text(section_data.get("body"))
        if section_title:
            document.add_heading(section_title, level=1)
        if section_body:
            _add_body_text(document, section_body)
        for chart_ref in _list_of_text(section_data.get("chart_refs")):
            inserted = _add_chart_for_ref(
                document=document,
                chart_ref=chart_ref,
                payload=payload,
                workspace_root=workspace_root,
                warnings=warnings,
            )
            if inserted:
                chart_asset_count += 1

    if payload.action_recommendations:
        document.add_heading("行动建议", level=1)
        for item in payload.action_recommendations:
            document.add_paragraph(item, style="List Bullet")

    if payload.data_boundaries:
        document.add_heading("数据边界", level=1)
        for item in payload.data_boundaries:
            document.add_paragraph(item, style="List Bullet")

    _add_evidence_appendix(document, payload)
    return chart_asset_count


def _configure_document_styles(document: Any) -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt, RGBColor

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12

    title = styles["Title"]
    title.font.name = "Microsoft YaHei"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(17, 24, 39)
    title.paragraph_format.space_after = Pt(8)

    for name, size in (("Heading 1", 14), ("Heading 2", 12)):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 41, 55)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)

    if "InsightFlowMeta" not in [style.name for style in styles]:
        meta = styles.add_style("InsightFlowMeta", WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta = styles["InsightFlowMeta"]
    meta.font.name = "Microsoft YaHei"
    meta._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    meta.font.size = Pt(9)
    meta.font.color.rgb = RGBColor(107, 114, 128)
    meta.paragraph_format.space_after = Pt(8)

    bullet = styles["List Bullet"]
    bullet.font.name = "Microsoft YaHei"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    bullet.font.size = Pt(10.5)
    bullet.paragraph_format.space_after = Pt(3)


def _add_body_text(document: Any, text: str) -> None:
    for paragraph_text in _paragraphs(text):
        document.add_paragraph(paragraph_text)


def _add_chart_for_ref(
    *,
    document: Any,
    chart_ref: str,
    payload: DocumentExportPayload,
    workspace_root: str | Path | None,
    warnings: list[str],
) -> bool:
    from docx.shared import Inches

    chart = payload.chart_artifacts.get(chart_ref) or {}
    asset = payload.static_assets.get(chart_ref) or _asset_for_chart(chart, payload.static_assets)
    title = _safe_text(chart.get("title")) or _safe_text(asset.get("title")) or chart_ref
    if title:
        document.add_heading(title, level=2)

    absolute_path = _asset_absolute_path(asset or chart, workspace_root=workspace_root)
    if absolute_path and absolute_path.exists() and absolute_path.suffix.lower() in SUPPORTED_WORD_IMAGE_EXTENSIONS:
        try:
            document.add_picture(str(absolute_path), width=Inches(5.9))
            annotation = _safe_text(chart.get("business_annotation"))
            if annotation:
                _add_body_text(document, annotation)
            return True
        except Exception:  # noqa: BLE001 - docx image parser errors become export warnings.
            warnings.append(f"图表 {chart_ref} 的静态图无法插入 Word，已写入占位说明。")

    reason = "缺少可插入的静态图资产"
    if absolute_path and absolute_path.suffix.lower() == ".svg":
        reason = "SVG 静态图已记录，但当前 Word 导出器需要 PNG/JPEG 等可嵌入图片"
    warnings.append(f"图表 {chart_ref} {reason}。")
    document.add_paragraph(f"图表暂无法插入：{title or chart_ref}。")
    return False


def _add_evidence_appendix(document: Any, payload: DocumentExportPayload) -> None:
    refs = _unique_texts([*payload.evidence_refs, *_list_of_text(payload.evidence_summary.get("refs"))])
    summary_items = []
    for key, label in (
        ("fact_count", "事实数"),
        ("table_count", "证据表数"),
        ("chart_count", "图表数"),
        ("ledger_fact_count", "账本事实数"),
        ("ledger_metric_count", "账本指标数"),
    ):
        value = payload.evidence_summary.get(key)
        if isinstance(value, int) and value > 0:
            summary_items.append(f"{label}：{value}")

    if not refs and not summary_items:
        return
    document.add_heading("证据附录", level=1)
    if summary_items:
        _add_body_text(document, "；".join(summary_items))
    if refs:
        for ref in refs[:40]:
            document.add_paragraph(ref, style="List Bullet")


def _payload_from_package(package: dict[str, Any]) -> DocumentExportPayload:
    document = package.get("document") if isinstance(package.get("document"), dict) else {}
    title = _safe_text(package.get("title")) or _safe_text(document.get("title")) or "报告"
    sections = []
    for section in package.get("sections") or document.get("sections") or []:
        if not isinstance(section, dict):
            continue
        safe_section = {
            "section_id": _safe_text(section.get("section_id")),
            "title": _safe_text(section.get("title")),
            "body": _safe_text(section.get("body")),
            "chart_refs": _list_of_text(section.get("chart_refs")),
            "evidence_refs": _list_of_text(section.get("evidence_refs")),
        }
        if safe_section["title"] or safe_section["body"]:
            sections.append(safe_section)
    return DocumentExportPayload(
        workspace_id=_safe_text(package.get("workspace_id")),
        source_id=_safe_text(package.get("source_id")),
        title=title,
        generated_at=_safe_text(package.get("generated_at")),
        summary=_safe_text(package.get("business_content_summary")) or _safe_text(document.get("opening_summary")),
        time_range=_safe_text(document.get("time_range")),
        data_sources=_list_of_text(document.get("data_sources")),
        sections=sections,
        action_recommendations=_list_of_text(package.get("action_recommendations")),
        data_boundaries=_list_of_text(package.get("data_boundaries")),
        chart_artifacts=_chart_index(package.get("chart_artifacts")),
        static_assets=_asset_index(package.get("static_assets")),
        evidence_refs=_list_of_text(package.get("evidence_refs")),
        evidence_summary=_safe_evidence_summary(package.get("evidence_summary")),
    )


def _chart_index(charts: Any) -> dict[str, dict[str, Any]]:
    indexed: dict[str, dict[str, Any]] = {}
    if not isinstance(charts, list):
        return indexed
    for chart in charts:
        if not isinstance(chart, dict):
            continue
        safe_chart = {
            "artifact_id": _safe_text(chart.get("artifact_id")),
            "chart_id": _safe_text(chart.get("chart_id")),
            "title": _safe_text(chart.get("title")),
            "path": _safe_relative_path(chart.get("path")),
            "image_path": _safe_relative_path(chart.get("image_path")),
            "url": _safe_text(chart.get("url")),
            "image_url": _safe_text(chart.get("image_url")),
            "business_annotation": _safe_text(chart.get("business_annotation")),
            "evidence_refs": _list_of_text(chart.get("evidence_refs")),
        }
        for key in (safe_chart["artifact_id"], safe_chart["chart_id"], safe_chart["title"]):
            if key:
                indexed[key] = safe_chart
    return indexed


def _asset_index(assets: Any) -> dict[str, dict[str, Any]]:
    indexed: dict[str, dict[str, Any]] = {}
    if not isinstance(assets, list):
        return indexed
    for asset in assets:
        if not isinstance(asset, dict):
            continue
        safe_asset = {
            "asset_id": _safe_text(asset.get("asset_id")),
            "title": _safe_text(asset.get("title")),
            "path": _safe_relative_path(asset.get("path")),
            "url": _safe_text(asset.get("url")),
            "format": _safe_text(asset.get("format")),
        }
        for key in (safe_asset["asset_id"], safe_asset["title"]):
            if key:
                indexed[key] = safe_asset
    return indexed


def _asset_for_chart(chart: dict[str, Any], assets: dict[str, dict[str, Any]]) -> dict[str, Any]:
    for key in (_safe_text(chart.get("artifact_id")), _safe_text(chart.get("chart_id")), _safe_text(chart.get("title"))):
        if key and key in assets:
            return assets[key]
    return {}


def _asset_absolute_path(asset: dict[str, Any], *, workspace_root: str | Path | None) -> Path | None:
    path = _safe_relative_path(asset.get("path") or asset.get("image_path"))
    if not path or not workspace_root:
        return None
    root = Path(workspace_root).resolve()
    candidate = (root / path).resolve()
    try:
        candidate.relative_to(root)
    except ValueError:
        return None
    return candidate


def _safe_evidence_summary(value: Any) -> dict[str, Any]:
    if not isinstance(value, dict):
        return {}
    summary: dict[str, Any] = {}
    for key in ("fact_count", "table_count", "chart_count", "ledger_fact_count", "ledger_metric_count"):
        if isinstance(value.get(key), int):
            summary[key] = value[key]
    for key in ("refs", "data_boundaries", "warnings", "data_limits", "notes", "chart_refs"):
        texts = _list_of_text(value.get(key))
        if texts:
            summary[key] = texts
    return summary


def _document_target(
    payload: DocumentExportPayload,
    *,
    workspace_root: str | Path | None,
    output_dir: str | Path | None,
) -> dict[str, Any]:
    root = Path(workspace_root).resolve() if workspace_root else Path.cwd().resolve()
    output = Path(str(output_dir or DEFAULT_DOCUMENT_EXPORT_DIR))
    if output.is_absolute() and not workspace_root:
        filename = f"{_safe_filename(payload.title)}_{_safe_filename(payload.source_id or 'report')}.docx"
        absolute_path = (output / filename).resolve()
        return {
            "absolute_path": absolute_path,
            "relative_path": absolute_path.as_posix(),
            "download_name": filename,
        }
    if output.is_absolute():
        try:
            relative_output = output.resolve().relative_to(root)
        except (OSError, ValueError):
            return {}
    else:
        relative_output = output
    if _unsafe_relative_path(relative_output.as_posix()):
        return {}
    filename = f"{_safe_filename(payload.title)}_{_safe_filename(payload.source_id or 'report')}.docx"
    relative_path = (relative_output / filename).as_posix()
    absolute_path = (root / relative_path).resolve()
    try:
        absolute_path.relative_to(root)
    except ValueError:
        return {}
    return {
        "absolute_path": absolute_path,
        "relative_path": relative_path,
        "download_name": filename,
    }


def _is_supported_report_package(package: dict[str, Any]) -> bool:
    return (
        isinstance(package, dict)
        and package.get("package_version", EXPORT_PACKAGE_VERSION) == EXPORT_PACKAGE_VERSION
        and package.get("source_type") == SOURCE_REPORT
    )


def _failed_result(warning: str) -> dict[str, Any]:
    return {
        "success": False,
        "document_path": "",
        "download_name": "",
        "warnings": [warning],
        "artifact": {},
    }


def _to_dict(value: Any) -> dict[str, Any]:
    if isinstance(value, dict):
        return value
    to_dict = getattr(value, "to_dict", None)
    if callable(to_dict):
        data = to_dict()
        return data if isinstance(data, dict) else {}
    return {}


def _paragraphs(text: str) -> list[str]:
    return [item.strip() for item in re.split(r"\n{1,}", str(text or "")) if item.strip()]


def _list_of_text(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    return [_safe_text(item) for item in value if _safe_text(item)]


def _unique_texts(values: list[str]) -> list[str]:
    return list(dict.fromkeys([_safe_text(value) for value in values if _safe_text(value)]))


def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    text = str(value).strip()
    return "" if _contains_sensitive_text(text) else text


def _safe_relative_path(value: Any) -> str:
    text = _safe_text(value)
    if not text:
        return ""
    if "://" in text or text.startswith("/api/"):
        return ""
    if _unsafe_relative_path(text):
        return ""
    path = Path(text)
    if path.is_absolute():
        return ""
    return path.as_posix()


def _unsafe_relative_path(value: str) -> bool:
    text = str(value or "").strip()
    if text.startswith("~"):
        return True
    return any(part == ".." for part in Path(text).parts)


def _contains_sensitive_text(value: str) -> bool:
    text = str(value or "").strip()
    if not text:
        return False
    lowered = text.lower()
    if re.search(r"\b(?:select|with|insert|update|delete|drop|alter|create|pragma)\b", text, re.IGNORECASE):
        return True
    if re.search(r"(^|[=/\s])(?:/users/|/tmp/|~[/\\])", lowered):
        return True
    if re.search(
        r"\b(?:raw_sql|generated_sql|raw_rows|trace_path|trace_id|provider_metadata|api_key|apikey|access_key|access_token|secret|password|database_path|analysis\.db|task_id|task_purpose|debug_id|prompt(?:_id|_text|_version)?|completion_tokens|prompt_tokens)\b",
        lowered,
    ):
        return True
    if re.search(r"\bsk-[A-Za-z0-9_-]+", text):
        return True
    parsed_path = urlsplit(text).path.lower()
    if lowered == "trace.json" or parsed_path.endswith((".db", "/trace.json")):
        return True
    return False


def _safe_filename(value: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9\u4e00-\u9fff_-]+", "_", str(value or "").strip())
    return safe.strip("_")[:80] or "report"


try:
    from docx.oxml.ns import qn
except ModuleNotFoundError:  # pragma: no cover - handled by export_report_docx dependency check.
    def qn(value: str) -> str:
        return value
